"""VLM-based document extraction for lab PDFs and intake forms via OpenRouter."""

from __future__ import annotations

import base64
import hashlib
import json
import logging
import re
import time
from typing import Any, Literal

import httpx
from pydantic import ValidationError

from app.schemas.extraction import IntakeFormResult, LabExtractionResult
from app.settings import Settings

DocType = Literal["lab_pdf", "intake_form"]

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_HL7_MIMES = frozenset({"application/hl7-v2", "application/hl7-v2+er7"})

_LAB_SIGNALS = (
    "reference range", "specimen", "abnormal", "loinc",
    "result value", "collected:", "lab report",
)
_INTAKE_SIGNALS = (
    "chief complaint", "chief concern", "patient information",
    "intake form", "current medications", "allergies",
    "family history", "date of birth",
)


def _heuristic_classify(text: str) -> DocType | None:
    t = text.lower()
    lab = sum(s in t for s in _LAB_SIGNALS)
    intake = sum(s in t for s in _INTAKE_SIGNALS)
    if lab >= 2 and lab > intake:
        return "lab_pdf"
    if intake >= 2 and intake > lab:
        return "intake_form"
    return None

_LOG = logging.getLogger("clinical_copilot.document_extractor")

# Approximate per-token USD cost for cost estimation (input_rate, output_rate).
_MODEL_COST_PER_TOKEN: dict[str, tuple[float, float]] = {
    "anthropic/claude-sonnet-4.6": (3.0e-6, 15.0e-6),
    "anthropic/claude-opus-4.7": (15.0e-6, 75.0e-6),
}
_DEFAULT_COST_PER_TOKEN: tuple[float, float] = (3.0e-6, 15.0e-6)


def _source_id(file_bytes: bytes) -> str:
    return "sha256:" + hashlib.sha256(file_bytes).hexdigest()[:16]


def _pdf_pages_to_images(file_bytes: bytes) -> list[tuple[str, str]]:
    """Render every PDF page to a base64 PNG. Returns list of (b64_data, mime_type)."""
    try:
        import fitz  # pymupdf
    except ModuleNotFoundError as exc:
        raise RuntimeError("PyMuPDF (fitz) is required for PDF extraction but is not installed.") from exc

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[tuple[str, str]] = []
    for page in doc:
        pix = page.get_pixmap(dpi=150)
        pages.append((base64.standard_b64encode(pix.tobytes("png")).decode(), "image/png"))
    return pages


def _tiff_pages_to_images(file_bytes: bytes) -> list[tuple[str, str]]:
    """Render every TIFF page to a base64 PNG. Returns list of (b64_data, mime_type)."""
    try:
        import fitz  # pymupdf
    except ModuleNotFoundError as exc:
        raise RuntimeError("PyMuPDF (fitz) is required for TIFF extraction but is not installed.") from exc

    doc = fitz.open(stream=file_bytes, filetype="tiff")
    pages: list[tuple[str, str]] = []
    for page in doc:
        pix = page.get_pixmap(dpi=150, alpha=False)
        pages.append((base64.standard_b64encode(pix.tobytes("png")).decode(), "image/png"))
    return pages


def _single_image_to_b64(file_bytes: bytes, mime_type: str) -> list[tuple[str, str]]:
    """Encode a single image file as base64. Returns a single-item list."""
    return [(base64.standard_b64encode(file_bytes).decode(), mime_type)]


def _docx_to_text(file_bytes: bytes) -> str:
    """Extract paragraph and table text from a DOCX document."""
    try:
        import io as _io
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX extraction.") from exc
    try:
        doc = Document(_io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(f"Cannot parse DOCX: {exc}") from exc
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx_to_text(file_bytes: bytes) -> str:
    """Extract cell values from all sheets of an XLSX workbook."""
    try:
        import io as _io
        import openpyxl  # type: ignore[import-untyped]
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for XLSX extraction.") from exc
    try:
        wb = openpyxl.load_workbook(_io.BytesIO(file_bytes), read_only=True, data_only=True)
    except Exception as exc:
        raise ValueError(f"Cannot parse XLSX: {exc}") from exc
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join("" if v is None else str(v) for v in row)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _hl7_to_text(file_bytes: bytes) -> str:
    """Extract segment-by-segment text from an HL7v2 message."""
    try:
        import hl7  # type: ignore[import-untyped]
    except ImportError as exc:
        raise RuntimeError("hl7 package is required for HL7v2 extraction.") from exc
    try:
        raw = file_bytes.decode("utf-8")
        normalized = raw.replace("\r\n", "\r").replace("\n", "\r")
        message = hl7.parse(normalized)
        return "\r".join(str(seg) for seg in message)
    except Exception as exc:
        raise ValueError(f"Cannot parse HL7v2 message: {exc}") from exc


def _text_parse_fallback(
    doc_type: DocType | None,
    sid: str,
    warning: str,
    start_time: float,
) -> tuple[LabExtractionResult | IntakeFormResult, dict[str, Any], int]:
    resolved: DocType = doc_type or "lab_pdf"
    fallback = _fallback_parsed_for_doc_type(resolved, warning)
    _inject_source_id(fallback, sid)
    validated: LabExtractionResult | IntakeFormResult
    if resolved == "lab_pdf":
        validated = LabExtractionResult.model_validate(fallback)
    else:
        validated = IntakeFormResult.model_validate(fallback)
    return validated, {}, int((time.perf_counter() - start_time) * 1000.0)


def _strip_markdown_fences(text: str) -> str:
    """Remove ```json … ``` fences so json.loads() can parse the body."""
    stripped = re.sub(r"^```(?:json)?\s*\n?", "", text.strip(), flags=re.IGNORECASE)
    return re.sub(r"\n?```\s*$", "", stripped).strip()


def _escape_control_chars_in_json_strings(raw: str) -> str:
    chars: list[str] = []
    in_string = False
    escaped = False
    for ch in raw:
        if escaped:
            chars.append(ch)
            escaped = False
            continue
        if ch == "\\":
            chars.append(ch)
            escaped = True
            continue
        if ch == '"':
            chars.append(ch)
            in_string = not in_string
            continue
        if in_string and ord(ch) < 0x20:
            chars.append(f"\\u{ord(ch):04x}")
            continue
        chars.append(ch)
    return "".join(chars)


def _extract_first_json_object(raw: str) -> str | None:
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    return raw[start:end + 1]


def _inject_source_id(parsed: dict[str, Any], sid: str) -> None:
    """Replace VLM source_id placeholders with the actual document hash."""
    for result in parsed.get("results", []):
        cit = result.get("citation")
        if isinstance(cit, dict):
            cit["source_id"] = sid
    cit = parsed.get("citation")
    if isinstance(cit, dict):
        cit["source_id"] = sid


def _json_schema_for_doc_type(doc_type: str) -> dict[str, Any]:
    if doc_type == "lab_pdf":
        return {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "schema_version": {"type": "string", "const": "1.0.0"},
                "doc_type": {"type": "string", "const": "lab_pdf"},
                "results": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "test_name": {"type": "string"},
                            "value": {"type": "string"},
                            "unit": {"type": "string"},
                            "reference_range": {"type": "string"},
                            "collection_date": {"type": "string"},
                            "abnormal_flag": {"type": "string"},
                            "confidence": {"type": "number", "minimum": 0.0, "maximum": 1.0},
                            "citation": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "source_type": {"type": "string"},
                                    "source_id": {"type": "string"},
                                    "page_or_section": {"type": "string"},
                                    "field_or_chunk_id": {"type": "string"},
                                    "quote_or_value": {"type": "string"},
                                    "bbox": {
                                        "anyOf": [
                                            {"type": "array", "items": {"type": "number"}, "minItems": 4, "maxItems": 4},
                                            {"type": "null"},
                                        ]
                                    },
                                    "page_number": {
                                        "anyOf": [{"type": "integer"}, {"type": "null"}]
                                    },
                                },
                                "required": [
                                    "source_type",
                                    "source_id",
                                    "page_or_section",
                                    "field_or_chunk_id",
                                    "quote_or_value",
                                    "bbox",
                                    "page_number",
                                ],
                            },
                        },
                        "required": [
                            "test_name",
                            "value",
                            "unit",
                            "reference_range",
                            "collection_date",
                            "abnormal_flag",
                            "confidence",
                            "citation",
                        ],
                    },
                },
                "extraction_warnings": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["schema_version", "doc_type", "results", "extraction_warnings"],
        }

    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "schema_version": {"type": "string", "const": "1.0.0"},
            "doc_type": {"type": "string", "const": "intake_form"},
            "demographics": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "name": {"type": "string"},
                    "dob": {"type": "string"},
                    "sex": {"type": "string"},
                    "address": {"type": "string"},
                },
                "required": ["name", "dob", "sex", "address"],
            },
            "chief_concern": {"type": "string"},
            "current_medications": {"type": "array", "items": {"type": "string"}},
            "allergies": {"type": "array", "items": {"type": "string"}},
            "family_history": {"type": "array", "items": {"type": "string"}},
            "extraction_warnings": {"type": "array", "items": {"type": "string"}},
            "citation": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "source_type": {"type": "string"},
                    "source_id": {"type": "string"},
                    "page_or_section": {"type": "string"},
                    "field_or_chunk_id": {"type": "string"},
                    "quote_or_value": {"type": "string"},
                    "bbox": {
                        "anyOf": [
                            {"type": "array", "items": {"type": "number"}, "minItems": 4, "maxItems": 4},
                            {"type": "null"},
                        ]
                    },
                    "page_number": {
                        "anyOf": [{"type": "integer"}, {"type": "null"}]
                    },
                },
                "required": [
                    "source_type",
                    "source_id",
                    "page_or_section",
                    "field_or_chunk_id",
                    "quote_or_value",
                    "bbox",
                    "page_number",
                ],
            },
        },
        "required": [
            "schema_version",
            "doc_type",
            "demographics",
            "chief_concern",
            "current_medications",
            "allergies",
            "family_history",
            "extraction_warnings",
            "citation",
        ],
    }


def _merge_usage(a: dict[str, Any], b: dict[str, Any]) -> dict[str, Any]:
    merged: dict[str, Any] = dict(a)
    for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
        merged[key] = int(merged.get(key, 0)) + int(b.get(key, 0))
    return merged


def _extract_message_content(resp_json: dict[str, Any]) -> str:
    content = resp_json["choices"][0]["message"]["content"]
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        chunks: list[str] = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text":
                text = item.get("text")
                if isinstance(text, str):
                    chunks.append(text)
        return "\n".join(chunks).strip()
    return str(content)


def _parse_json_with_fallbacks(raw_content: str) -> dict[str, Any] | None:
    candidates: list[str] = []
    stripped = _strip_markdown_fences(raw_content)
    candidates.append(raw_content)
    candidates.append(stripped)
    candidates.append(_escape_control_chars_in_json_strings(stripped))
    extracted = _extract_first_json_object(stripped)
    if extracted is not None:
        candidates.append(extracted)
        candidates.append(_escape_control_chars_in_json_strings(extracted))

    for candidate in candidates:
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            continue
    return None


def _fallback_parsed_for_doc_type(doc_type: str, warning: str) -> dict[str, Any]:
    if doc_type == "lab_pdf":
        return {
            "schema_version": "1.0.0",
            "doc_type": "lab_pdf",
            "results": [],
            "extraction_warnings": [warning],
        }
    return {
        "schema_version": "1.0.0",
        "doc_type": "intake_form",
        "demographics": {"name": "", "dob": "", "sex": "", "address": ""},
        "chief_concern": "",
        "current_medications": [],
        "allergies": [],
        "family_history": [],
        "extraction_warnings": [warning],
        "citation": {
            "source_type": "intake_form",
            "source_id": "PLACEHOLDER",
            "page_or_section": "unknown",
            "field_or_chunk_id": "intake_form_summary",
            "quote_or_value": "Unable to parse model output reliably.",
        },
    }


async def _call_openrouter_completion(
    client: httpx.AsyncClient,
    *,
    settings: Settings,
    model: str,
    content: list[dict[str, Any]],
    doc_type: str,
    enforce_schema: bool,
) -> tuple[dict[str, Any], dict[str, Any], str]:
    headers = {
        "Authorization": f"Bearer {settings.openrouter_api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": settings.openrouter_http_referer,
        "X-Title": settings.openrouter_app_title,
    }
    payload: dict[str, Any] = {
        "model": model,
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 4096,
    }
    if enforce_schema:
        payload["response_format"] = {
            "type": "json_schema",
            "json_schema": {
                "name": f"{doc_type}_extraction",
                "strict": True,
                "schema": _json_schema_for_doc_type(doc_type),
            },
        }

    response = await client.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers=headers,
        json=payload,
    )
    response.raise_for_status()
    resp_json: dict[str, Any] = response.json()
    token_usage: dict[str, Any] = resp_json.get("usage", {})
    return resp_json, token_usage, _extract_message_content(resp_json)


async def _repair_json_once(
    client: httpx.AsyncClient,
    *,
    settings: Settings,
    model: str,
    doc_type: str,
    raw_content: str,
) -> tuple[dict[str, Any] | None, dict[str, Any]]:
    repair_prompt = (
        "Fix the following model output into valid JSON that matches the requested extraction schema exactly. "
        "Return JSON only, no prose, no markdown.\n\n"
        f"doc_type={doc_type}\n\n"
        "REMINDER: Every citation object MUST include \"bbox\" ([x0,y0,x1,y1] floats or null) "
        "and \"page_number\" (1-indexed integer or null).\n\n"
        "RAW OUTPUT START\n"
        + raw_content[:8000]
        + "\nRAW OUTPUT END"
    )
    repair_content: list[dict[str, Any]] = [{"type": "text", "text": repair_prompt}]
    try:
        resp_json, token_usage, repaired = await _call_openrouter_completion(
            client,
            settings=settings,
            model=model,
            content=repair_content,
            doc_type=doc_type,
            enforce_schema=True,
        )
        _ = resp_json
        return _parse_json_with_fallbacks(repaired), token_usage
    except Exception:
        return None, {}


def estimate_cost_usd(model: str, token_usage: dict[str, Any]) -> float:
    input_rate, output_rate = _MODEL_COST_PER_TOKEN.get(model, _DEFAULT_COST_PER_TOKEN)
    return round(
        int(token_usage.get("prompt_tokens", 0)) * input_rate
        + int(token_usage.get("completion_tokens", 0)) * output_rate,
        6,
    )


def _lab_pdf_prompt() -> str:
    return (
        "You are a medical document AI. Extract ALL laboratory results from the document image(s).\n\n"
        "Respond with ONLY valid JSON — no markdown fences, no prose — matching this schema:\n"
        '{"schema_version":"1.0.0","doc_type":"lab_pdf","results":[{"test_name":"<string>",'
        '"value":"<raw printed value>","unit":"<string or empty>","reference_range":"<string or empty>",'
        '"collection_date":"<ISO-8601 or empty>","abnormal_flag":"<H|L|HH|LL|A or empty>",'
        '"confidence":<0.0-1.0>,"citation":{"source_type":"lab_pdf","source_id":"PLACEHOLDER",'
        '"page_or_section":"<page N>","field_or_chunk_id":"<test slug>",'
        '"quote_or_value":"<verbatim text>","bbox":[x0,y0,x1,y1],"page_number":<N>}}],'
        '"extraction_warnings":[]}\n\n'
        "Rules: extract every analyte visible; copy verbatim text into quote_or_value; "
        "lower confidence for unclear scans; use empty string for absent fields. "
        "Each citation MUST include \"bbox\" as [x0,y0,x1,y1] in PDF points (origin top-left) "
        "and \"page_number\" (1-indexed integer) when the source location is determinable; "
        "set both to null for HL7/DOCX/XLSX sources where spatial coordinates are not meaningful."
    )


def _intake_form_prompt() -> str:
    return (
        "You are a medical document AI. Extract patient information from the intake form image(s).\n\n"
        "Respond with ONLY valid JSON — no markdown fences, no prose — matching this schema:\n"
        '{"schema_version":"1.0.0","doc_type":"intake_form",'
        '"demographics":{"name":"","dob":"","sex":"","address":""},'
        '"chief_concern":"","current_medications":[],"allergies":[],"family_history":[],'
        '"extraction_warnings":[],'
        '"citation":{"source_type":"intake_form","source_id":"PLACEHOLDER",'
        '"page_or_section":"page 1","field_or_chunk_id":"intake_form_summary",'
        '"quote_or_value":"<brief verbatim excerpt>","bbox":[x0,y0,x1,y1],"page_number":<N>}}\n\n'
        "Rules: empty string for absent text fields; empty list for absent list fields; "
        "include medication name and dose as written; note illegible text in extraction_warnings. "
        "The citation MUST include \"bbox\" as [x0,y0,x1,y1] in PDF points (origin top-left) "
        "and \"page_number\" (1-indexed integer) when the source location is determinable; "
        "set both to null for HL7/DOCX/XLSX sources where spatial coordinates are not meaningful."
    )


async def classify_doc_type(
    *,
    image_payload: list[bytes] | None = None,
    text_content: str | None = None,
    settings: Settings,
) -> DocType:
    """Classify a document as lab_pdf or intake_form.

    Strategy:
    - If text_content available, try heuristic first (zero cost).
    - Otherwise (or on heuristic abstain), make ONE VLM call with the first
      page only, asking for a single token: "lab_pdf" or "intake_form".
    - Default to "lab_pdf" only if the VLM returns garbage; emit a warning.
    """
    if text_content is not None:
        result = _heuristic_classify(text_content[:2048])
        if result is not None:
            return result

    if not image_payload:
        _LOG.warning("classify_doc_type_no_content defaulting to lab_pdf")
        return "lab_pdf"

    classify_prompt = (
        'You are classifying a clinical document. Reply with EXACTLY one word, '
        'either "lab_pdf" or "intake_form". No punctuation. No explanation.'
    )
    first_page_b64 = base64.standard_b64encode(image_payload[0]).decode()
    content: list[dict[str, Any]] = [
        {"type": "text", "text": classify_prompt},
        {
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{first_page_b64}", "detail": "low"},
        },
    ]

    try:
        async with httpx.AsyncClient(timeout=settings.openrouter_http_timeout_s) as client:
            _, _, raw = await _call_openrouter_completion(
                client,
                settings=settings,
                model=settings.vlm_model,
                content=content,
                doc_type="lab_pdf",
                enforce_schema=False,
            )
        normalized = raw.strip().lower().rstrip(".,!;: ")
        if "intake" in normalized:
            return "intake_form"
        if "lab" in normalized:
            return "lab_pdf"
        _LOG.warning("classify_doc_type_vlm_unexpected_response raw=%s defaulting to lab_pdf", raw[:100])
        return "lab_pdf"
    except Exception:
        _LOG.warning("classify_doc_type_vlm_failed defaulting to lab_pdf", exc_info=True)
        return "lab_pdf"


async def extract_document(
    file_bytes: bytes,
    mime_type: str,
    doc_type: DocType | None,
    settings: Settings,
) -> tuple[LabExtractionResult | IntakeFormResult, dict[str, Any], int]:
    """Extract structured data from a document via VLM.

    Returns (validated_result, token_usage_dict, latency_ms).
    Returns safe partial structured output when parsing or strict validation fails.
    Raises ValueError for unsupported or unrecognised MIME types.
    """
    req_start = time.perf_counter()
    sid = _source_id(file_bytes)

    # Phase 1: Prepare content — images for PDF/TIFF/images, text for office/HL7 formats.
    image_payload: list[tuple[str, str]] | None = None
    extracted_text: str | None = None

    if mime_type == "application/pdf":
        image_payload = _pdf_pages_to_images(file_bytes)
    elif mime_type in ("image/tiff", "image/x-tiff"):
        image_payload = _tiff_pages_to_images(file_bytes)
    elif mime_type == _DOCX_MIME:
        try:
            extracted_text = _docx_to_text(file_bytes)
        except ValueError as exc:
            _LOG.warning("document_extractor_docx_parse_error error=%s", exc)
            return _text_parse_fallback(doc_type, sid, f"Failed to parse DOCX: {exc}", req_start)
    elif mime_type == _XLSX_MIME:
        try:
            extracted_text = _xlsx_to_text(file_bytes)
        except ValueError as exc:
            _LOG.warning("document_extractor_xlsx_parse_error error=%s", exc)
            return _text_parse_fallback(doc_type, sid, f"Failed to parse XLSX: {exc}", req_start)
    elif mime_type in _HL7_MIMES:
        try:
            extracted_text = _hl7_to_text(file_bytes)
        except ValueError as exc:
            _LOG.warning("document_extractor_hl7_parse_error error=%s", exc)
            return _text_parse_fallback(doc_type, sid, f"Failed to parse HL7v2: {exc}", req_start)
    elif mime_type == "text/plain":
        raw_text = file_bytes.decode("utf-8", errors="replace")
        first_nonempty = next((line for line in raw_text.splitlines() if line.strip()), "")
        if first_nonempty.startswith("MSH|"):
            try:
                extracted_text = _hl7_to_text(file_bytes)
            except ValueError as exc:
                _LOG.warning("document_extractor_hl7_parse_error error=%s", exc)
                return _text_parse_fallback(doc_type, sid, f"Failed to parse HL7v2: {exc}", req_start)
        else:
            raise ValueError("text/plain without HL7 MSH| prefix is not supported")
    elif mime_type == "application/msword":
        raise ValueError(
            "Legacy .doc format (application/msword) is not supported; convert to .docx first"
        )
    elif mime_type.startswith("image/"):
        image_payload = _single_image_to_b64(file_bytes, mime_type)
    else:
        raise ValueError(f"Unsupported MIME type: {mime_type!r}")

    # Phase 2: Classify doc type if not provided by caller.
    if doc_type is None:
        if extracted_text is not None:
            doc_type = await classify_doc_type(text_content=extracted_text, settings=settings)
        elif image_payload is not None:
            first_page_b64, _ = image_payload[0]
            first_page_bytes = base64.standard_b64decode(first_page_b64)
            doc_type = await classify_doc_type(
                image_payload=[first_page_bytes],
                settings=settings,
            )
        else:
            doc_type = "lab_pdf"

    # Phase 3: Build VLM content list (text-only path or image path).
    prompt = _lab_pdf_prompt() if doc_type == "lab_pdf" else _intake_form_prompt()
    content: list[dict[str, Any]]
    if extracted_text is not None:
        content = [{"type": "text", "text": f"{prompt}\n\n{extracted_text}"}]
        page_count = 0
    else:
        assert image_payload is not None
        content = [{"type": "text", "text": prompt}]
        for b64_data, img_mime in image_payload:
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{img_mime};base64,{b64_data}", "detail": "high"},
            })
        page_count = len(image_payload)

    # Phase 4: Call VLM, parse, repair if needed, validate.
    model = settings.vlm_model
    token_usage: dict[str, Any] = {}
    async with httpx.AsyncClient(timeout=settings.openrouter_http_timeout_s) as client:
        raw_content = ""
        try:
            _, call_usage, raw_content = await _call_openrouter_completion(
                client,
                settings=settings,
                model=model,
                content=content,
                doc_type=doc_type,
                enforce_schema=True,
            )
            token_usage = _merge_usage(token_usage, call_usage)
        except httpx.HTTPStatusError as exc:
            if exc.response is None or exc.response.status_code not in (400, 422):
                response_body = exc.response.text if exc.response is not None else ""
                if len(response_body) > 2000:
                    response_body = response_body[:2000] + "...<truncated>"
                _LOG.error(
                    "document_extractor_openrouter_http_error status=%d model=%s doc_type=%s body=%s",
                    exc.response.status_code if exc.response is not None else 0,
                    model,
                    doc_type,
                    response_body,
                )
                raise

            _LOG.warning(
                "document_extractor_schema_mode_not_supported status=%d model=%s doc_type=%s",
                exc.response.status_code,
                model,
                doc_type,
            )
            _, call_usage, raw_content = await _call_openrouter_completion(
                client,
                settings=settings,
                model=model,
                content=content,
                doc_type=doc_type,
                enforce_schema=False,
            )
            token_usage = _merge_usage(token_usage, call_usage)

        parsed = _parse_json_with_fallbacks(raw_content)
        if parsed is None:
            repaired, repair_usage = await _repair_json_once(
                client,
                settings=settings,
                model=model,
                doc_type=doc_type,
                raw_content=raw_content,
            )
            token_usage = _merge_usage(token_usage, repair_usage)
            parsed = repaired

        if parsed is None:
            output_snippet = raw_content
            if len(output_snippet) > 2000:
                output_snippet = output_snippet[:2000] + "...<truncated>"
            _LOG.warning(
                "document_extractor_parse_error doc_type=%s model=%s output=%s",
                doc_type,
                model,
                output_snippet,
            )
            parsed = _fallback_parsed_for_doc_type(
                doc_type,
                "Document was difficult to parse; returning safe partial structured result.",
            )

    _inject_source_id(parsed, sid)

    result: LabExtractionResult | IntakeFormResult
    try:
        if doc_type == "lab_pdf":
            result = LabExtractionResult.model_validate(parsed)
        else:
            result = IntakeFormResult.model_validate(parsed)
    except ValidationError as exc:
        _LOG.warning(
            "document_extractor_schema_validation_error doc_type=%s model=%s error=%s",
            doc_type,
            model,
            str(exc).replace("\n", " ")[:2000],
        )
        fallback = _fallback_parsed_for_doc_type(
            doc_type,
            "Extracted content did not match schema strictly; returning safe partial structured result.",
        )
        _inject_source_id(fallback, sid)
        if doc_type == "lab_pdf":
            result = LabExtractionResult.model_validate(fallback)
        else:
            result = IntakeFormResult.model_validate(fallback)

    latency_ms = int((time.perf_counter() - req_start) * 1000.0)
    _LOG.info(
        "document_extractor_ok doc_type=%s pages=%d model=%s latency_ms=%d prompt_tokens=%d completion_tokens=%d",
        doc_type,
        page_count,
        model,
        latency_ms,
        int(token_usage.get("prompt_tokens", 0)),
        int(token_usage.get("completion_tokens", 0)),
    )

    return result, token_usage, latency_ms
